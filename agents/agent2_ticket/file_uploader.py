import os
import json
from io import BytesIO
from typing import List, Dict, Optional
from pathlib import Path

# File processing libraries
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

from langchain_community.embeddings import OllamaEmbeddings
from qdrant_client import QdrantClient
from qdrant_client.models import PointStruct


class FileProcessor:
    """Process various file types and extract text content"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.extract_pdf,
            '.docx': self.extract_docx,
            '.doc': self.extract_docx,
            '.txt': self.extract_txt,
            '.md': self.extract_txt,
            '.json': self.extract_json,
            '.xlsx': self.extract_excel,
            '.xls': self.extract_excel,
        }
    
    def is_supported(self, filename: str) -> bool:
        """Check if file format is supported"""
        ext = Path(filename).suffix.lower()
        return ext in self.supported_formats
    
    def extract_text(self, file_path: str = None, file_content: bytes = None, filename: str = None) -> str:
        """Extract text from file"""
        if filename:
            ext = Path(filename).suffix.lower()
        elif file_path:
            ext = Path(file_path).suffix.lower()
        else:
            raise ValueError("Either filename or file_path must be provided")
        
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {ext}")
        
        extractor = self.supported_formats[ext]
        
        if file_content:
            return extractor(file_content=file_content)
        elif file_path:
            return extractor(file_path=file_path)
        else:
            raise ValueError("Either file_path or file_content must be provided")
    
    def extract_pdf(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from PDF"""
        if not pdfplumber and not PyPDF2:
            raise ImportError("PDF processing requires pypdf2 or pdfplumber")
        
        text = ""
        
        # Try pdfplumber first (better text extraction)
        if pdfplumber:
            try:
                if file_content:
                    pdf_file = BytesIO(file_content)
                    with pdfplumber.open(pdf_file) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                else:
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                return text.strip()
            except Exception as e:
                print(f"pdfplumber failed: {e}, trying PyPDF2")
        
        # Fallback to PyPDF2
        if PyPDF2:
            try:
                if file_content:
                    pdf_file = BytesIO(file_content)
                    reader = PyPDF2.PdfReader(pdf_file)
                else:
                    reader = PyPDF2.PdfReader(file_path)
                
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
            except Exception as e:
                raise Exception(f"Failed to extract PDF: {e}")
        
        raise Exception("No PDF processor available")
    
    def extract_docx(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from Word document"""
        if not Document:
            raise ImportError("Word processing requires python-docx")
        
        try:
            if file_content:
                doc_file = BytesIO(file_content)
                doc = Document(doc_file)
            else:
                doc = Document(file_path)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract Word document: {e}")
    
    def extract_txt(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from plain text or markdown file"""
        try:
            if file_content:
                return file_content.decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
        except Exception as e:
            raise Exception(f"Failed to extract text file: {e}")
    
    def extract_json(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from JSON file"""
        try:
            if file_content:
                data = json.loads(file_content.decode('utf-8'))
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
            
            # Convert JSON to readable text
            return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            raise Exception(f"Failed to extract JSON: {e}")
    
    def extract_excel(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from Excel file"""
        if not openpyxl:
            raise ImportError("Excel processing requires openpyxl")
        
        try:
            if file_content:
                excel_file = BytesIO(file_content)
                wb = openpyxl.load_workbook(excel_file)
            else:
                wb = openpyxl.load_workbook(file_path)
            
            text = ""
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"\n=== Sheet: {sheet_name} ===\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    text += row_text + "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract Excel: {e}")


class DocumentUploader:
    """Upload documents to Qdrant vector database"""
    
    def __init__(
        self,
        qdrant_url: str = "http://qdrant:6333",
        collection_name: str = "agent2_tickets",
        embeddings_model: str = "llama3",
        ollama_url: str = "http://ollama:11434",
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ):
        self.qdrant_client = QdrantClient(url=qdrant_url)
        self.collection_name = collection_name
        self.embeddings = OllamaEmbeddings(model=embeddings_model, base_url=ollama_url)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.file_processor = FileProcessor()
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        chunks = []
        words = text.split()
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk = " ".join(words[i:i + self.chunk_size])
            if chunk.strip():
                chunks.append(chunk)
        
        return chunks
    
    def upload_file(
        self,
        file_path: str = None,
        file_content: bytes = None,
        filename: str = None,
        category: str = "Document",
        metadata: Dict = None
    ) -> Dict:
        """
        Upload a file to Qdrant
        
        Args:
            file_path: Path to the file
            file_content: File content as bytes
            filename: Name of the file
            category: Category for the document
            metadata: Additional metadata
        
        Returns:
            Dict with upload results
        """
        try:
            # Extract text from file
            print(f"📄 Processing file: {filename or file_path}")
            text = self.file_processor.extract_text(
                file_path=file_path,
                file_content=file_content,
                filename=filename or file_path
            )
            
            if not text.strip():
                return {"error": "No text content found in file"}
            
            print(f"✅ Extracted {len(text)} characters from file")
            
            # Chunk the text
            chunks = self.chunk_text(text)
            
            if not chunks:
                return {"error": "No chunks created from text"}
            
            print(f"📦 Created {len(chunks)} chunks from text")
            
            # Create points for Qdrant
            points = []
            point_id = hash(filename or file_path) % (10**8)
            
            for i, chunk in enumerate(chunks):
                try:
                    # Generate embedding
                    embedding = self.embeddings.embed_query(chunk)
                    
                    # Prepare metadata
                    payload = {
                        "text": chunk,
                        "category": category,
                        "type": "document",
                        "filename": filename or Path(file_path).name,
                        "chunk_index": i,
                        "total_chunks": len(chunks)
                    }
                    
                    # Add custom metadata
                    if metadata:
                        payload.update(metadata)
                    
                    # Create point
                    point = PointStruct(
                        id=point_id + i,
                        vector=embedding,
                        payload=payload
                    )
                    points.append(point)
                    print(f"  ⚡ Generated embedding for chunk {i+1}/{len(chunks)}")
                    
                except Exception as e:
                    print(f"❌ Error processing chunk {i}: {e}")
                    continue
            
            if not points:
                return {"error": "No points created from chunks"}
            
            # Upload to Qdrant
            print(f"⬆️  Uploading {len(points)} points to Qdrant collection '{self.collection_name}'...")
            self.qdrant_client.upsert(
                collection_name=self.collection_name,
                points=points
            )
            
            print(f"✅ Successfully uploaded '{filename or Path(file_path).name}' to Qdrant!")
            print(f"   📊 {len(points)} chunks | 📁 Category: {category}")
            
            return {
                "status": "success",
                "filename": filename or Path(file_path).name,
                "chunks_uploaded": len(points),
                "category": category,
                "collection": self.collection_name
            }
            
        except Exception as e:
            return {"error": str(e)}
    
    def upload_directory(
        self,
        directory_path: str,
        category: str = "Document",
        recursive: bool = True
    ) -> List[Dict]:
        """Upload all supported files from a directory"""
        results = []
        
        path = Path(directory_path)
        
        if recursive:
            files = path.rglob("*")
        else:
            files = path.glob("*")
        
        for file_path in files:
            if file_path.is_file() and self.file_processor.is_supported(file_path.name):
                print(f"Processing: {file_path}")
                result = self.upload_file(
                    file_path=str(file_path),
                    filename=file_path.name,
                    category=category
                )
                results.append(result)
        
        return results
    
    def get_collection_stats(self) -> Dict:
        """Get statistics about the collection"""
        try:
            collection_info = self.qdrant_client.get_collection(self.collection_name)
            
            # Count document types
            scroll_result = self.qdrant_client.scroll(
                collection_name=self.collection_name,
                limit=1000,
                with_payload=True
            )
            
            doc_types = {}
            categories = {}
            
            for point in scroll_result[0]:
                payload = point.payload
                doc_type = payload.get("type", "unknown")
                category = payload.get("category", "unknown")
                
                doc_types[doc_type] = doc_types.get(doc_type, 0) + 1
                categories[category] = categories.get(category, 0) + 1
            
            return {
                "collection_name": self.collection_name,
                "total_points": collection_info.points_count,
                "vector_size": collection_info.config.params.vectors.size,
                "document_types": doc_types,
                "categories": categories
            }
        except Exception as e:
            return {"error": str(e)}


# Convenience functions
def upload_file_to_qdrant(
    file_path: str,
    category: str = "Document",
    collection_name: str = "agent2_tickets",
    qdrant_url: str = "http://qdrant:6333",
    ollama_url: str = "http://ollama:11434"
) -> Dict:
    """Quick function to upload a single file"""
    uploader = DocumentUploader(
        qdrant_url=qdrant_url,
        collection_name=collection_name,
        ollama_url=ollama_url
    )
    return uploader.upload_file(file_path=file_path, category=category)


def upload_directory_to_qdrant(
    directory_path: str,
    category: str = "Document",
    collection_name: str = "agent2_tickets",
    qdrant_url: str = "http://qdrant:6333",
    ollama_url: str = "http://ollama:11434",
    recursive: bool = True
) -> List[Dict]:
    """Quick function to upload all files from a directory"""
    uploader = DocumentUploader(
        qdrant_url=qdrant_url,
        collection_name=collection_name,
        ollama_url=ollama_url
    )
    return uploader.upload_directory(directory_path, category, recursive)
